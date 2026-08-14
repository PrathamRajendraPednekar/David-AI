import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor

def create_report():
    doc = Document()
    
    # Title
    title = doc.add_heading('David AI - Architectural & Technical Report', 0)
    title.alignment = 1 # Center
    
    # ---------------- OVERVIEW ---------------- #
    doc.add_heading('1. Executive Overview', level=1)
    doc.add_paragraph(
        "David AI is an autonomous, voice-centric intelligent desktop assistant capable of orchestrating complex "
        "tasks autonomously. Unlike standard conversational chatbots, David fundamentally acts as a desktop agent— "
        "writing files independently, opening computer applications, dispatching texts and routing complex video calls "
        "over WhatsApp through native PC control. It employs a multi-threaded architecture with cross-compatible AI "
        "inference engines to bypass server downtimes seamlessly."
    )
    
    # ---------------- APIS ---------------- #
    doc.add_heading('2. Application Programming Interfaces (APIs)', level=1)
    
    doc.add_paragraph("The engine relies on cloud infrastructure to parse logic, translate languages, and retrieve current data.", style="List Bullet")
    
    apis = [
        ("Groq API Cloud", "Operates as the Primary Priority AI layer. Uses the `llama-3.3-70b-versatile` model. Capable of incredibly fast (instantaneous inference) data parsing. Handles context extraction, generative conversation, and logical phone number identification."),
        ("Google Gemini AI API", "Serves as the Secondary Priority Intelligence engine (`gemini-2.5-flash`). Seamlessly configured as an active failsafe. If Groq hits rate-limit ceilings, David silently dynamically switches back to Google Gemini, guaranteeing a zero-downtime execution."),
        ("SerpApi (Google Search Engine)", "Allows the assistant to silently execute localized or global Google searches during questions (such as live weather, recent news, or stock queries) before handing the context directly to Groq."),
        ("HuggingFace Inference API", "Powering the Image Generation layer. Takes creative structural text and synthesizes visual files natively to the user's local disk.")
    ]
    for title, desc in apis:
        p = doc.add_paragraph()
        run = p.add_run(f"• {title}: ")
        run.bold = True
        p.add_run(desc)
        
    # ---------------- LIBRARIES ---------------- #
    doc.add_heading('3. Python Libraries & Ecosystem Data', level=1)
    doc.add_paragraph("A comprehensive network of localized Python modules dictate the agent's mechanical and operational abilities.", style="List Bullet")
    
    libs = [
        ("pyautogui", "Used for Mechanical GUI Automation. Simulates human mouse movements, dragging hover animations, and executes specialized keyboard clicks to physically interface with the user's WhatsApp desktop ecosystem autonomously."),
        ("opencv-python", "Used in tandem with `pyautogui` for specialized local Computer Vision functionality. David 'looks' at the current desktop layout, searches for visual neural-matches (like the WhatsApp 'Call' icon or a 'Video Call' button), validates the pixel coordinates against false positives, and executes the physical click!"),
        ("pywhatkit", "Utilized in parsing and directly routing user messages logically via deep-links or browsers into WhatsApp contacts efficiently."),
        ("groq & google-genai", "The native SDKs connecting David's logic loops physically to the external LLM cloud interfaces."),
        ("edge-tts", "Employed for Advanced Text-to-Speech logic. Operates asynchronously from the main logic thread, allowing David to articulate his current plan softly alongside task execution without locking up processing cores."),
        ("speech_recognition", "The core driver capturing the user's raw microphone microphone feed and utilizing dynamic noise-threshold calibration to understand intent seamlessly."),
        ("PyQt5", "Renders the dynamic, multi-threaded frontend GUI layout. Controls the sleek window configurations, internal interactive widgets (like the custom microphone/send icons), and fluid context switching when messages require loading."),
        ("python-docx", "Used to construct this precise document algorithmically via Python!")
    ]
    for title, desc in libs:
        p = doc.add_paragraph()
        run = p.add_run(f"{title}")
        run.bold = True
        run.font.color.rgb = RGBColor(0, 51, 153)
        p.add_run(f": {desc}")
        
    # ---------------- AUTOMATION ---------------- #
    doc.add_heading('4. Advanced Automation Routines', level=1)
    
    auto_desc = doc.add_paragraph()
    auto_desc.add_run("Dynamic Document Generation:\n").bold = True
    auto_desc.add_run("David listens to trigger words including spelling typos ('emil', 'appilcation', 'leave application'). These logic flags immediately divert standard Chat algorithms. David isolates the LLM's raw text, automatically dumps it into a '.txt' file (`generated_content.txt` or `leave_application.txt`) into his local `Content AI/` storage, and dynamically executes `notepad.exe` using standard PC subprocesses so the resulting code simply opens instantly on the user's desktop view.")
    
    auto_desc2 = doc.add_paragraph()
    auto_desc2.add_run("\nComputer Vision Deep Routing (WhatsApp Media Calls):\n").bold = True
    auto_desc2.add_run("David intercepts 'Call [Contact]' requests securely. He deploys LLM extractions (Priority 1 Groq -> Priority 2 Gemini) to grab the specific phone number. He deploys browser URI deep-links (`whatsapp://send`) to snap the WhatsApp specific chat open. Next, he relies strictly on `opencv-python` to scan the physical Desktop layout with 85-90% AI certainty vectors against user-supplied reference images. By capturing X,Y coordinate meshes, he smoothly sweeps the mouse toward the call dropdowns to simulate human movement and securely avoids clicking identical false interfaces.")
    
    # Save the document
    report_path = os.path.join(os.path.dirname(__file__), "David_AI_Project_Report_V2.docx")
    doc.save(report_path)
    print(f"Report fully compiled & generated at: {report_path}")
    
    # Open the document
    os.startfile(report_path)

if __name__ == "__main__":
    create_report()
